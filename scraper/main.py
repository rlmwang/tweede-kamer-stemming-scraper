# A new beginning, let the behaviour be known
import os
import re
import ssl
from collections.abc import Iterator
from io import BytesIO
from pathlib import Path

import click
import polars as pl
import requests
import wget
from bs4 import BeautifulSoup
from docx import Document

RESULTS_ROOT = Path("results")

STEMMINGSUITSLAGEN_URL = (
    "https://www.tweedekamer.nl/kamerstukken/stemmingsuitslagen"
    "?qry=%2A&fld_tk_categorie=Kamerstukken&fld_prl_kamerstuk=Stemmingsuitslagen"
    "&srt=date%3Adesc%3Adate&page={page}"
)
DEBAT_URL = "https://www.tweedekamer.nl/{link}"
MOTIE_URL = "https://www.tweedekamer.nl/{link}"
DOWNLOAD_URL = "https://www.tweedekamer.nl/{link}"

STEMMING_SCHEMA = {
    "stemming_id": str,
    "stemming_did": str,
    "titel": str,
    "datum": str,
    "type": str,
}
MOTIE_SCHEMA = {
    "stemming_id": str,
    "motie_id": str,
    "motie_did": str,
    "document_nr": str,
    "datum": str,
    "titel": str,
    "type": str,
    "text": str,
    "is_fallback": bool,
    "download": str,
    "besluit": str,
    "uitslag": str,
    "voor": int,
    "vereist": int,
    "totaal": int,
}
INDIENERS_SCHEMA = {
    "motie_id": str,
    "name": str,
    "type": str,
}
DETAILS_SCHEMA = {
    "motie_id": str,
    "fractie": str,
    "zetels": str,
    "stem": str,
}


def run(begin_page: int, end_page: int | None = None):
    end_page = max(begin_page or 0, end_page or 0)

    result = None
    for i in range(begin_page, end_page + 1):
        url = STEMMINGSUITSLAGEN_URL.format(page=i)
        for data in parse_listings_page(url=url):
            if len(data["stemming"]) != 1:
                raise ValueError(f"Multiple votings in one page for {url}")
            stemming_id = data["stemming"]["stemming_id"].item()

            write_tables(data, stemming_id)


def parse_listings_page(url) -> Iterator[dict[str, pl.DataFrame]]:
    resp = requests.get(url)
    if not resp.ok:
        raise ValueError(f"Page {url} does not respond")

    page = resp.content
    soup = BeautifulSoup(page, "lxml")
    links = [a["href"] for a in soup.select("h4.u-mt-0 > a")]

    if len(links) == 0:
        raise ValueError(f"No links found for {url}")

    result = None
    for link in links:
        print(link)
        yield parse_stemming_page(url=DEBAT_URL.format(link=link.strip("/")))


def parse_stemming_page(url) -> dict[str, pl.DataFrame]:
    data = create_tables()

    resp = requests.get(url)
    if not resp.ok:
        raise ValueError(f"Page {url} does not respond")
    page = resp.content
    soup = BeautifulSoup(page, "lxml")

    # parse voting

    stemming_info = parse_stemming_page_info(url=url, soup=soup)
    data["stemming"] = pl.concat([data["stemming"], pl.DataFrame(stemming_info)])

    # parse individual motions

    cards = soup.select("div.m-card")
    if len(cards) == 0:
        raise ValueError(f"No motions found for {url}")

    for k, card in enumerate(cards):
        # Find the main motion link
        link_tag = card.select_one("h3.m-card__title > a")
        if link_tag is None:
            raise ValueError(f"Cannot find link of motion {k} for {url}")
        link = link_tag["href"]
        print("  " + link)

        # Find the paragraph with "Besluit"
        besluit_tag = None
        for p in card.select("p.u-mt-8"):
            if "Besluit" in p.get_text():
                besluit_tag = p.select_one("span.u-font-bold")
                break
        if besluit_tag is None:
            raise ValueError(f"Cannot find decision of motion {k} for {url}")
        besluit = besluit_tag.get_text(strip=True).strip(".")

        motie_data = parse_motie_page(
            url=MOTIE_URL.format(link=link.strip("/")),
            stemming_id=stemming_info["stemming_id"],
            besluit=besluit,
        )

        data = merge_tables(data, motie_data)

    return data


def parse_stemming_page_info(url: str, soup) -> dict:
    # stemming id & did
    match = re.search(r"id=([^&]+)&did=([^&]+)", url)
    if not match:
        raise ValueError(f"ID and DID missing from motion for {url}")
    stemming_id = match.group(1)
    stemming_did = match.group(2)

    # stemming title
    meta_tag = soup.select_one('meta[name="dcterms.title"]')
    if meta_tag is None:
        ValueError(f"Stemming title is missing for {url}")
    stemming_title = meta_tag["content"]

    # stemming type & date
    h2_tags = soup.select("h2")
    if len(h2_tags) != 1:
        raise ValueError(f"Stemming subtitle is missing or not unique for {url}")

    for h2 in h2_tags:
        span = h2.select_one("span.u-font-normal")
        if span is None:
            raise ValueError(f"Stemming date is missing for {url}")
        stemming_date = span.get_text(strip=True)

        # The meeting type is everything in h2 minus the span
        stemming_type = h2.contents[0].strip() if h2.contents else None
        if type is None:
            raise ValueError(f"Stemming type is missing for {url}")

    return {
        "stemming_id": stemming_id,
        "stemming_did": stemming_did,
        "titel": stemming_title,
        "datum": stemming_date,
        "type": stemming_type,
    }


def parse_motie_page(
    url: str,
    stemming_id: str,
    besluit: str | None,
) -> dict[str, pl.DataFrame]:
    data = create_tables()

    resp = requests.get(url)
    if not resp.ok:
        raise ValueError(f"Page {url} does not respond")
    page = resp.content
    soup = BeautifulSoup(page, "lxml")

    motie_info = {
        "stemming_id": stemming_id,
        **parse_motie_info(url=url, soup=soup),
        "besluit": besluit,
    }

    indieners_info = parse_indieners_info(url=url, soup=soup)
    for row in indieners_info:
        row["motie_id"] = motie_info["motie_id"]

    h2 = soup.find("h2", string=lambda t: t and "Stemmingsuitslagen" in t)
    if h2 is not None:
        # motie uitslag text
        h3 = h2.find_next("h3")
        if h3 is None:
            raise ValueError(f"Uitslag is missing from stemmingsuitslagen for {url}")
        motie_info["uitslag"] = h3.get_text(strip=True)

        # motie uitslag counts
        labels = soup.select("div.m-vote-result__label")
        motie_info["voor"] = int(labels[0].find("span").get_text(strip=True).lstrip(": "))
        motie_info["vereist"] = int(labels[1].find("span").get_text(strip=True).split(": ")[1])
        motie_info["totaal"] = int(labels[2].find("span").get_text(strip=True).split(": ")[1])

        # motie uitslag details
        rows = soup.select("#votes-details table.h-table-bordered tbody tr")[1:]
        details_info = [
            {
                "fractie": r.find_all("td")[0].get_text(strip=True),
                "zetels": int(r.find_all("td")[1].get_text(strip=True)),
                "stem": r.find_all("td")[2].get_text(strip=True),
            }
            for r in rows
        ]
        for row in details_info:
            row["motie_id"] = motie_info["motie_id"]
    else:
        for key in ["uitslag", "voor", "vereist", "totaal"]:
            motie_info[key] = None
        details_info = []

    # dump data into tables
    data["motie"] = pl.concat([data["motie"], pl.DataFrame(motie_info, schema=MOTIE_SCHEMA)])
    data["indieners"] = pl.concat(
        [data["indieners"], pl.DataFrame(indieners_info, schema=INDIENERS_SCHEMA)]
    )
    data["details"] = pl.concat(
        [data["details"], pl.DataFrame(details_info, schema=DETAILS_SCHEMA)]
    )

    return data


def parse_motie_info(url: str, soup) -> dict:
    # motie id & did
    match = re.search(r"id=([^&]+)&did=([^&]+)", url)
    if not match:
        raise ValueError(f"ID and DID missing from motion for {url}")
    motie_id = match.group(1)
    motie_did = match.group(2)

    # motie document nr
    doc_nr_spans = soup.select("span.h-visually-hidden")
    all_doc_nrs = [
        span.next_sibling.strip() for span in doc_nr_spans if "Nummer:" in span.get_text()
    ]
    if len(all_doc_nrs) != 1:
        raise ValueError(f"Document nr missing or not unique for {url}")
    motie_document_nr = all_doc_nrs[0]

    # motie date
    date_spans = soup.select("span.h-visually-hidden")
    all_dates = [span.next_sibling.strip() for span in date_spans if "Datum:" in span.get_text()]
    if len(all_dates) != 1:
        raise ValueError(f"Date missing or not unique for {url}")
    motie_date = all_dates[0]

    # motion type & title
    all_titles = []
    for h1 in soup.select("h1"):
        span = h1.select_one("span.u-text-primary.u-font-normal")
        if span:
            motion_type = span.get_text(strip=True)
            # collect all non-span text inside the h1
            title_text = "".join(t for t in h1.stripped_strings if t != motion_type and t != ":")
            all_titles.append({"type": motion_type, "titel": title_text})

    # motie pdf url
    download_tag = soup.select_one('a[aria-label^="Download kamerstuk"]')
    if download_tag is None:
        raise ValueError(f"Motie PDF is missing for {url}")
    motie_download = DOWNLOAD_URL.format(link=download_tag["href"].strip("/"))

    # motion text
    content = soup.select_one("div.m-modal__content")

    if content:
        is_fallback = False

        # normalize whitespace
        motie_text = " ".join(t.strip() for t in content.stripped_strings)
        motie_text = " ".join(motie_text.split())
    else:
        is_fallback = True

        # fallback to parsing download
        response = requests.get(motie_download)
        response.raise_for_status()
        doc = Document(BytesIO(response.content))

        text_parts = []

        # paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text.strip())

        # tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())

        # combine and normalize whitespace
        motie_text = " ".join(text_parts)
        motie_text = " ".join(motie_text.split())

    if len(all_titles) != 1:
        raise ValueError(f"Title missing or not unique for {url}")
    motie_type, motie_title = all_titles[0]["type"], all_titles[0]["titel"]

    return {
        "motie_id": motie_id,
        "motie_did": motie_did,
        "document_nr": motie_document_nr,
        "datum": motie_date,
        "titel": motie_title,
        "type": motie_type,
        "text": motie_text,
        "is_fallback": is_fallback,
        "download": motie_download,
    }


def parse_indieners_info(url: str, soup) -> list[dict]:
    indieners = []
    for li in soup.select("ul.m-list li.m-list__item--variant-member"):
        # indiener type

        type_span = li.select_one("span.u-font-bold")
        if type_span is None:
            raise ValueError(f"Type of (mede)indiener is missing for {url}")
        type_text = type_span.get_text(strip=True)

        # indiener name
        label_span = li.select_one("span.m-list__label")
        if not label_span:
            raise ValueError(f"Name of (mede)indiener is missing for {url}")

        name_link = label_span.select_one("a.h-link-inverse")
        if name_link:
            # Try first <a>
            name_text = name_link.get_text(strip=True)
        else:
            # Take the text after the type span
            texts = [t.strip() for t in label_span.stripped_strings]
            if len(texts) <= 1:
                raise ValueError(f"Name of (mede)indiener is missing for {url}")
            full_name = texts[1]

            # Split off descriptor at first comma
            if "," in full_name:
                *name_parts, _ = [s.strip() for s in full_name.split(",")]
                name_text = ", ".join(part.strip() for part in name_parts)
            else:
                name_text = full_name

        indieners.append({"type": type_text, "name": name_text})

    return indieners


# UTILS


def create_tables() -> dict[str, pl.DataFrame]:
    return {
        "stemming": pl.DataFrame(schema=STEMMING_SCHEMA),
        "motie": pl.DataFrame(schema=MOTIE_SCHEMA),
        "indieners": pl.DataFrame(schema=INDIENERS_SCHEMA),
        "details": pl.DataFrame(schema=DETAILS_SCHEMA),
    }


def write_tables(data: dict[str, pl.DataFrame], name: str):
    (RESULTS_ROOT / name).mkdir(parents=True, exist_ok=True)
    for key, table in data.items():
        table.write_csv(RESULTS_ROOT / name / f"{key}.csv")


def merge_tables(
    data_a: dict[str, pl.DataFrame] | None,
    data_b: dict[str, pl.DataFrame],
) -> dict[str, pl.DataFrame]:
    if data_a is None:
        return data_b
    output = {}
    for key in data_a.keys():
        output[key] = pl.concat([data_a[key], data_b[key]])
    return output


@click.command()
@click.argument("begin_page", type=int)
@click.argument("end_page", type=int, required=False)
def cli(begin_page, end_page):
    """Scrape Tweede Kamer motions from BEGIN_PAGE to END_PAGE."""
    run(begin_page, end_page)


if __name__ == "__main__":
    cli()
