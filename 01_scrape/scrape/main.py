# A new beginning, let the behaviour be known
import json
import os
import re
import ssl
from collections.abc import Iterator
from datetime import date, datetime
from io import BytesIO
from pathlib import Path

import click
import magic
import polars as pl
import requests
import wget
from bs4 import BeautifulSoup
from dateparser import parse as parse_date
from docx import Document
from PyPDF2 import PdfReader

STEMMINGSUITSLAGEN_URL = (
    "https://www.tweedekamer.nl/kamerstukken/stemmingsuitslagen"
    "?qry=%2A&fld_tk_categorie=Kamerstukken&fld_prl_kamerstuk=Stemmingsuitslagen"
    "&fromdate={from_date}&todate={to_date}"
    "&srt=date%3Adesc%3Adate&page={page}"
)
DEBAT_URL = "https://www.tweedekamer.nl/{link}"
MOTIE_URL = "https://www.tweedekamer.nl/{link}"
DOWNLOAD_URL = "https://www.tweedekamer.nl/{link}"

STEMMING_SCHEMA = {
    "stemming_id": str,
    "stemming_did": str,
    "titel": str,
    "datum": str,
    "type": str,
}
MOTIE_SCHEMA = {
    "stemming_id": str,
    "motie_id": str,
    "motie_did": str,
    "document_nr": str,
    "datum": str,
    "titel": str,
    "type": str,
    "text": str,
    "is_fallback": bool,
    "download": str,
    "besluit": str,
    "uitslag": str,
    "voor": int,
    "vereist": int,
    "totaal": int,
}
INDIENERS_SCHEMA = {
    "stemming_id": str,
    "motie_id": str,
    "name": str,
    "type": str,
}
DETAILS_SCHEMA = {
    "stemming_id": str,
    "motie_id": str,
    "fractie": str,
    "zetels": str,
    "kamerlid": str,
    "stem": str,
    "niet_deelgenomen": str,
    "vergissing": bool,
}


def run(
    from_date: date,
    to_date: date | None,
    output_dir: str,
    full_refresh: bool,
    select: str | None,
):
    to_date = max(to_date or date.today(), from_date)
    select = parse_select_argument(select)

    progress = read_progress()

    page = 0
    result = None
    while True:
        print(f"Page {page:02d}")

        url = STEMMINGSUITSLAGEN_URL.format(from_date=from_date, to_date=to_date, page=page)
        resp = requests.get(url)
        if not resp.ok:
            raise ValueError(f"Page {url} does not respond")

        if "Geen zoekresultaten" in resp.text:
            print(f"No further pages found.")
            break

        soup = BeautifulSoup(resp.content, "lxml")
        for data in parse_listings_page(
            url=url, soup=soup, select=select, progress=progress, full_refresh=full_refresh
        ):
            if len(data["stemming"]) != 1:
                raise ValueError(f"Multiple votings in one page for {url}")

            stem_id = data["stemming"]["stemming_id"].item()
            stem_dt = data["stemming"]["datum"].item()
            stem_dt = parse_dutch_date_str(stem_dt)

            write_tables(data, Path(output_dir) / stem_dt / stem_id)

        page += 1


def parse_listings_page(
    url: str,
    soup: BeautifulSoup,
    select: list[str] | None,
    progress: dict[str, list],
    full_refresh: bool,
) -> Iterator[dict[str, pl.DataFrame]]:
    cards = []
    for card in soup.select("div.m-card, div.u-mt-6.m-card"):
        a_tag = card.select_one("h4.u-mt-0 > a")
        if a_tag:
            link = a_tag["href"]

        time_tag = card.select_one("time.u-text-primary")
        if time_tag:
            stem_dt = time_tag.get_text(strip=True)
            stem_dt = parse_dutch_date_str(stem_dt)

        id_tag = card.select_one("p.u-text-dark-gray")
        if id_tag:
            stem_id = id_tag.get_text(strip=True)

        cards.append(
            {
                "link": link,
                "stem_dt": stem_dt,
                "stem_id": stem_id,
            }
        )

    if len(cards) == 0:
        raise ValueError(f"No links found for {url}")

    for card in cards:
        if select is not None and card["stem_id"] not in select:
            continue

        if not full_refresh and already_processed(progress, card["stem_dt"], card["stem_id"]):
            continue

        print(card["stem_dt"], card["stem_id"], card["link"])
        result, ok = parse_stemming_page(url=DEBAT_URL.format(link=card["link"].strip("/")))
        yield result

        if ok:
            rem_error(stem_id=card["stem_id"])

        progress.setdefault(card["stem_dt"], []).append(card["stem_id"])
        write_progress(progress)


def parse_stemming_page(url: str) -> tuple[dict[str, pl.DataFrame], bool]:
    result = create_tables()
    res_ok = True

    resp = requests.get(url)
    if not resp.ok:
        raise ValueError(f"Page {url} does not respond")
    soup = BeautifulSoup(resp.content, "lxml")

    # parse voting

    stemming_info = parse_stemming_page_info(url=url, soup=soup)
    result["stemming"] = pl.concat([result["stemming"], pl.DataFrame(stemming_info)])

    # parse individual motions

    cards = soup.select("div.m-card")
    if len(cards) == 0:
        return result, res_ok

    for k, card in enumerate(cards):
        # Find the main motion link
        link_tag = card.select_one("h3.m-card__title > a")
        if link_tag is None:
            raise ValueError(f"Cannot find link of motion {k} for {url}")
        link = link_tag["href"]
        print("  " + link)

        # Find the paragraph with "Besluit"
        besluit_tag = None
        for p in card.select("p.u-mt-8"):
            if "Besluit" in p.get_text():
                besluit_tag = p.select_one("span.u-font-bold")
                break
        if besluit_tag is None:
            raise ValueError(f"Cannot find decision of motion {k} for {url}")
        besluit = besluit_tag.get_text(strip=True).strip(".")

        try:
            motie_data = parse_motie_page(
                url=MOTIE_URL.format(link=link.strip("/")),
                stemming_id=stemming_info["stemming_id"],
                besluit=besluit,
            )
        except Exception as err:
            print(f"Failed {url}")
            rem_error(
                stem_id=stemming_info["stemming_id"],
            )
            add_error(
                stem_id=stemming_info["stemming_id"],
                url=MOTIE_URL.format(link=link.strip("/")),
                err=err,
            )
            res_ok = False
            continue

        result = merge_tables(result, motie_data)

        rem_error(
            stem_id=stemming_info["stemming_id"],
            url=MOTIE_URL.format(link=link.strip("/")),
        )

    return result, res_ok


def parse_stemming_page_info(url: str, soup: BeautifulSoup) -> dict:
    # stemming id & did
    match = re.search(r"id=([^&]+)&(?:did|dossier)=([^&]+)", url)
    if not match:
        raise ValueError(f"ID and DID missing from motion for {url}")
    stemming_id = match.group(1)
    stemming_did = match.group(2)

    # stemming title
    meta_tag = soup.select_one('meta[name="dcterms.title"]')
    if meta_tag is None:
        ValueError(f"Stemming title is missing for {url}")
    stemming_title = meta_tag["content"]

    # stemming type & date
    h2_tags = soup.select("h2")
    if len(h2_tags) != 1:
        raise ValueError(f"Stemming subtitle is missing or not unique for {url}")

    for h2 in h2_tags:
        span = h2.select_one("span.u-font-normal")
        if span is None:
            raise ValueError(f"Stemming date is missing for {url}")
        stemming_date = span.get_text(strip=True)

        # The meeting type is everything in h2 minus the span
        stemming_type = h2.contents[0].strip() if h2.contents else None
        if type is None:
            raise ValueError(f"Stemming type is missing for {url}")

    return {
        "stemming_id": stemming_id,
        "stemming_did": stemming_did,
        "titel": stemming_title,
        "datum": stemming_date,
        "type": stemming_type,
    }


def parse_motie_page(
    url: str,
    stemming_id: str,
    besluit: str | None,
) -> dict[str, pl.DataFrame]:
    data = create_tables()

    resp = requests.get(url)
    if not resp.ok:
        raise ValueError(f"Page {url} does not respond")
    soup = BeautifulSoup(resp.content, "lxml")

    motie_info = parse_motie_info(url=url, soup=soup)
    motie_info = {
        "stemming_id": stemming_id,
        **motie_info,
        "besluit": besluit,
    }

    indieners_info = parse_indieners_info(url=url, soup=soup)
    for row in indieners_info:
        row["stemming_id"] = motie_info["stemming_id"]
        row["motie_id"] = motie_info["motie_id"]

    h2 = soup.find("h2", string=lambda t: t and "Stemmingsuitslagen" in t)
    if h2 is not None:
        # motie uitslag text
        h3 = h2.find_next("h3")
        if h3 is None:
            raise ValueError(f"Uitslag is missing from stemmingsuitslagen for {url}")
        motie_info["uitslag"] = h3.get_text(strip=True)

        # motie uitslag counts
        labels = soup.select("div.m-vote-result__label")
        motie_info["voor"] = int(labels[0].find("span").get_text(strip=True).lstrip(": "))
        motie_info["vereist"] = int(labels[1].find("span").get_text(strip=True).split(": ")[1])
        motie_info["totaal"] = int(labels[2].find("span").get_text(strip=True).split(": ")[1])

        # motie uitslag details
        details_info = parse_details_info(url=url, soup=soup)
        for row in details_info:
            row["stemming_id"] = motie_info["stemming_id"]
            row["motie_id"] = motie_info["motie_id"]
    else:
        for key in ["uitslag", "voor", "vereist", "totaal"]:
            motie_info[key] = None
        details_info = []

    # dump data into tables
    data["motie"] = pl.concat([data["motie"], pl.DataFrame(motie_info, schema=MOTIE_SCHEMA)])
    data["indieners"] = pl.concat(
        [data["indieners"], pl.DataFrame(indieners_info, schema=INDIENERS_SCHEMA)]
    )
    data["details"] = pl.concat(
        [data["details"], pl.DataFrame(details_info, schema=DETAILS_SCHEMA)]
    )

    return data


def parse_motie_info(url: str, soup: BeautifulSoup) -> dict:
    # motie id & did
    match = re.search(r"id=([^&]+)&(?:did|dossier)=([^&]+)", url)
    if not match:
        raise ValueError(f"ID and DID missing from motion for {url}")
    motie_id = match.group(1)
    motie_did = match.group(2)

    # motie document nr
    doc_nr_spans = soup.select("span.h-visually-hidden")
    all_doc_nrs = [
        span.next_sibling.strip() for span in doc_nr_spans if "Nummer:" in span.get_text()
    ]
    if len(all_doc_nrs) != 1:
        raise ValueError(f"Document nr missing or not unique for {url}")
    motie_document_nr = all_doc_nrs[0]

    # motie date
    date_spans = soup.select("span.h-visually-hidden")
    all_dates = [span.next_sibling.strip() for span in date_spans if "Datum:" in span.get_text()]
    if len(all_dates) != 1:
        raise ValueError(f"Date missing or not unique for {url}")
    motie_date = all_dates[0]

    # motion type & title
    all_titles = []
    for h1 in soup.select("h1"):
        span = h1.select_one("span.u-text-primary.u-font-normal")
        if span:
            motion_type = span.get_text(strip=True)
            # collect all non-span text inside the h1
            title_text = "".join(t for t in h1.stripped_strings if t != motion_type and t != ":")
            all_titles.append({"type": motion_type, "titel": title_text})

    if len(all_titles) != 1:
        raise ValueError(f"Title missing or not unique for {url}")
    motie_type, motie_title = all_titles[0]["type"], all_titles[0]["titel"]

    # check for early exits
    if motie_type.lower() in ["wetsvoorstel", "voorstel van wet", "eindtekst"]:
        # TODO: implement wetsvoorstellen
        return {
            "motie_id": motie_id,
            "motie_did": motie_did,
            "document_nr": motie_document_nr,
            "datum": motie_date,
            "titel": motie_title,
            "type": motie_type,
            "text": None,
            "is_fallback": False,
            "download": None,
        }

    # motie pdf url
    download_tag = soup.select_one('a[aria-label^="Download kamerstuk"]')
    if download_tag is None:
        raise ValueError(f"Motie PDF is missing for {url}")
    motie_download = DOWNLOAD_URL.format(link=download_tag["href"].strip("/"))

    # motion text
    content = soup.select_one("div.m-modal__content")
    if content:
        # normalize whitespace
        is_fallback = False
        motie_text = " ".join(t.strip() for t in content.stripped_strings)
        motie_text = " ".join(motie_text.split())
    else:
        # fallback to parsing download
        is_fallback = True
        motie_text = parse_text_from_download(url=motie_download)

    return {
        "motie_id": motie_id,
        "motie_did": motie_did,
        "document_nr": motie_document_nr,
        "datum": motie_date,
        "titel": motie_title,
        "type": motie_type,
        "text": motie_text,
        "is_fallback": is_fallback,
        "download": motie_download,
    }


def parse_text_from_download(url: str) -> str:
    # download the file
    response = requests.get(url)
    response.raise_for_status()
    file_bytes = BytesIO(response.content)

    # detect file type
    file_type = magic.from_buffer(file_bytes.getvalue(), mime=True)

    text_parts = []
    if "wordprocessingml" in file_type:
        # DOCX parser
        doc = Document(file_bytes)

        # paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text.strip())

        # tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())

    elif file_type == "application/pdf":
        # PDF parser
        reader = PdfReader(file_bytes)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text.strip())

    else:
        raise ValueError(f"Unsupported file type: {file_type}")

    # combine and normalize whitespace
    motie_text = " ".join(text_parts)
    motie_text = " ".join(motie_text.split())

    return motie_text


EXPECTED_HEADERS = [
    ["Fracties", "Zetels", "Voor/Tegen"],
    ["Fracties", "Zetels", "Kamerlid", "Voor/Tegen"],
]


def parse_details_info(url: str, soup: BeautifulSoup) -> list[dict]:
    rows = soup.select("#votes-details table.h-table-bordered tbody tr")[1:]  # skip header row

    headers = [th.get_text(strip=True) for th in soup.select("#votes-details table thead th")]
    if not headers:
        headers = [
            th.get_text(strip=True)
            for th in soup.select("#votes-details table tbody tr")[0].find_all(["th", "td"])
        ]

    if headers[:3] != EXPECTED_HEADERS[0] and headers[:4] != EXPECTED_HEADERS[1]:
        raise ValueError(f"Unexpected table headers: {headers}")

    current_fractie = None
    current_zetels = None

    details_info = []
    for r in rows:
        cells = r.find_all(["td", "th"])

        if len(cells) == len(headers):
            # first row of a bloc with rowspan
            current_fractie = cells[0].get_text(strip=True)
            current_zetels = int(cells[1].get_text(strip=True))
            row = {
                h.lower().replace(" ", "_"): (c.get_text(strip=True) or None)
                for h, c in zip(headers[2:], cells[2:])
            }
        else:
            # subsequent rows
            row = {
                h.lower().replace(" ", "_"): (c.get_text(strip=True) or None)
                for h, c in zip(headers[2:], cells)
            }

        details_info.append(
            {
                "fractie": current_fractie,
                "zetels": current_zetels,
                "kamerlid": row.get("kamerlid"),
                "stem": row.get("voor/tegen"),
                "niet_deelgenomen": row.get("niet_deelgenomen"),
                "vergissing": bool(row.get("vergissing")),
            }
        )

    return details_info


def parse_indieners_info(url: str, soup: BeautifulSoup) -> list[dict]:
    indieners = []
    for li in soup.select("ul.m-list li.m-list__item--variant-member"):
        # indiener type

        type_span = li.select_one("span.u-font-bold")
        if type_span is None:
            raise ValueError(f"Type of (mede)indiener is missing for {url}")
        type_text = type_span.get_text(strip=True)

        # indiener name
        label_span = li.select_one("span.m-list__label")
        if not label_span:
            raise ValueError(f"2: Name of (mede)indiener is missing for {url}")

        name_link = label_span.select_one("a.h-link-inverse")
        if name_link:
            # Try first <a>
            name_text = name_link.get_text(strip=True)
        else:
            # Take the text after the type span
            texts = [t.strip() for t in label_span.stripped_strings]
            if len(texts) <= 1:
                continue
            full_name = texts[1]

            # Split off descriptor at first comma
            if "," in full_name:
                *name_parts, _ = [s.strip() for s in full_name.split(",")]
                name_text = ", ".join(part.strip() for part in name_parts)
            else:
                name_text = full_name

        indieners.append({"type": type_text, "name": name_text})

    if len(indieners) == 0:
        raise ValueError(f"(Mede)indieners are missing for {url}")

    return indieners


# UTILS


def parse_select_argument(arg: str | None) -> list[str] | None:
    return arg.split() if arg is not None else None


def create_tables() -> dict[str, pl.DataFrame]:
    return {
        "stemming": pl.DataFrame(schema=STEMMING_SCHEMA),
        "motie": pl.DataFrame(schema=MOTIE_SCHEMA),
        "indieners": pl.DataFrame(schema=INDIENERS_SCHEMA),
        "details": pl.DataFrame(schema=DETAILS_SCHEMA),
    }


def write_tables(data: dict[str, pl.DataFrame], path: Path):
    path.mkdir(parents=True, exist_ok=True)
    for key, table in data.items():
        table.write_csv(path / f"{key}.csv")


def merge_tables(
    data_a: dict[str, pl.DataFrame] | None,
    data_b: dict[str, pl.DataFrame],
) -> dict[str, pl.DataFrame]:
    if data_a is None:
        return data_b
    output = {}
    for key in data_a.keys():
        output[key] = pl.concat([data_a[key], data_b[key]])
    return output


def read_error() -> pl.DataFrame:
    file_path = Path(".run") / "errors.csv"
    if file_path.exists():
        return pl.read_csv(file_path)
    else:
        return pl.DataFrame(schema={"stemming_id": str, "url": str, "error": str})


def write_error(err_data: pl.DataFrame):
    file_path = Path(".run") / "errors.csv"
    file_path.parent.mkdir(parents=True, exist_ok=True)
    err_data = err_data.unique().sort("stemming_id", "url")
    err_data.write_csv(file_path)


def add_error(stem_id: str, url: str, err: Exception):
    err_row = pl.DataFrame(
        {
            "stemming_id": stem_id,
            "url": url,
            "error": str(err),
        }
    )
    err_data = read_error()
    err_data = pl.concat([err_data, err_row])
    write_error(err_data)


def rem_error(stem_id: str, url: str | None = None):
    err_data = read_error()
    if url is None:
        err_data = err_data.filter(
            pl.col("stemming_id") != stem_id,
        )
    else:
        err_data = err_data.filter(
            ~((pl.col("stemming_id") == stem_id) & (pl.col("url") == url)),
        )
    write_error(err_data)


def read_progress() -> dict[str, list]:
    file_path = Path(".run") / "progress.json"
    if file_path.exists():
        with file_path.open(encoding="utf-8") as f:
            progress = json.load(f)
    else:
        progress = {}
    progress = remove_failed_from_progress(progress)
    return progress


def remove_failed_from_progress(progress: dict[str, list]) -> dict[str, list]:
    err_data = read_error()
    err_ids = err_data["stemming_id"].to_list()

    result = {}
    for date, ids in progress.items():
        result[date] = [i for i in ids if i not in err_ids]

    return result


def write_progress(progress: dict[str, list]):
    file_path = Path(".run") / "progress.json"
    file_path.parent.mkdir(parents=True, exist_ok=True)
    with open(file_path, "w", encoding="utf-8") as f:
        json.dump(progress, f, indent=2, ensure_ascii=False)


def rebuild_progress(data_path: str):
    data_path = Path(data_path)
    file_path = Path(".run") / "progress.json"

    res = {}
    for folder in data_path.iterdir():
        if not folder.is_dir():
            continue
        key = folder.name
        res[key] = set()
        for subfolder in folder.iterdir():
            if not subfolder.is_dir():
                continue
            res[key].add(subfolder.name)
        res[key] = list(res[key])
    write_progress(res)


def already_processed(progress: dict[str, list], stem_dt: str, stem_id: str) -> bool:
    return stem_dt in progress and stem_id in progress[stem_dt]


def parse_dutch_date_str(date: str) -> str:
    date = parse_date(date, languages=["nl"])
    date = date.strftime("%Y-%m-%d")
    return date
